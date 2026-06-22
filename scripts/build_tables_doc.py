# -*- coding: utf-8 -*-
"""
build_tables_doc.py — מייצר מסמך Word בעברית שמסביר את טבלאות הבסיס של הפרויקט.
מסמך הסבר (לא חלק מצינור הנתונים) — נכתב לבקשת פאדי.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

AZURE = RGBColor(0x0B, 0x6B, 0x9C)
DARK = RGBColor(0x10, 0x20, 0x33)


def set_rtl(paragraph):
    """הופך פסקה לכיווניות ימין-לשמאל (עברית)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def rtl_run(run):
    """מסמן ריצת טקסט כ-RTL כדי שהעברית תוצג נכון."""
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def style_doc(doc):
    """מגדיר גופן בסיסי שתומך בעברית."""
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), "Arial")


def add_par(doc, text="", size=11, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    set_rtl(p)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        rtl_run(r)
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 20, 1: 15, 2: 12.5}
    p = add_par(doc, text, size=sizes.get(level, 12), bold=True,
                color=AZURE if level else DARK, space_after=8)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    r = p.add_run(text)
    rtl_run(r)
    r.font.size = Pt(size)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    """מוסיף טבלה עם כותרת מודגשת, בכיווניות ימין-לשמאל."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # כיווניות RTL לטבלה (העמודה הראשונה מימין)
    tblPr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], "0B6B9C")
        para = hdr[i].paragraphs[0]
        set_rtl(para)
        run = para.add_run(h)
        rtl_run(run)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for c_i, val in enumerate(row):
            if c_i < len(headers):
                shade_cell(cells[c_i], "EAF4FA" if r_i % 2 == 0 else "FFFFFF")
                para = cells[c_i].paragraphs[0]
                set_rtl(para)
                run = para.add_run(str(val))
                rtl_run(run)
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    doc.add_paragraph()
    return table


# ===========================================================================
def build():
    doc = Document()
    style_doc(doc)

    # ----- כותרת -----
    add_heading(doc, "הסבר מורחב על טבלאות הבסיס", level=0)
    add_par(doc, "פרויקט סקאוטינג כדורגל · Master Scout · פאדי והב · 2026",
            size=12, bold=True, color=AZURE)
    add_par(doc, "מסמך זה מסביר בעברית פשוטה מאילו נתונים יצאנו לדרך, ממה כל טבלה "
                 "הורכבה, ואילו עיבודים (\"מניפולציות\") ביצענו עליהן — עם דגש על "
                 "טבלת האירועים, שבה הושקעה רוב העבודה המקדימה הקשה.", size=11)

    # ----- 1. נקודת המוצא -----
    add_heading(doc, "1. נקודת המוצא — שתי טבלאות גולמיות", level=1)
    add_par(doc, "הורדנו שני מאגרים ציבוריים מ-Kaggle. אלה טבלאות \"גולמיות\" — בדיוק "
                 "כפי שהגיעו, לפני כל ניקוי:")
    add_table(doc,
              ["המאגר", "הקובץ", "שורות", "עמודות", "מה הוא מכיל"],
              [["EA Sports FC 24", "male_players.csv", "180,021", "109",
                "פרופיל שחקנים. 10 גרסאות FIFA (15–24)."],
               ["Football Events", "events.csv", "941,009", "22",
                "אירועי משחק בודדים, ליגות אירופה ~2012–2017."]],
              col_widths=[1.4, 1.5, 0.8, 0.7, 2.4])
    add_par(doc, "מתוך מאגר FC24 שמרנו רק את גרסת FC24 (fifa_version == 24) — כך כל "
                 "שורה = שחקן אחד, ונשארנו עם 18,350 שחקנים. מאגר האירועים מתאר "
                 "9,074 משחקים, 6,118 שחקנים ו-24,446 גולים.", bold=False)

    # ----- 2. טבלת השחקנים -----
    add_heading(doc, "2. טבלת השחקנים (FC24) — ממה הורכבה", level=1)
    add_par(doc, "כל שורה היא שחקן אחד. מתוך 109 העמודות בחרנו את הרלוונטיות בלבד, "
                 "מחולקות לקבוצות:")
    add_table(doc,
              ["קבוצת עמודות", "דוגמאות", "למה שמרנו"],
              [["זהות ומטא", "short_name, age, nationality, club, league, preferred_foot",
                "לזיהוי השחקן ולסינון לפי גיל/לאום/ליגה."],
               ["שוק וערך", "value_eur, wage_eur, overall, potential",
                "הבסיס לחיפוש לפי דירוג ולמציאת \"מציאות\"."],
               ["פיזי", "height_cm, weight_kg, weak_foot, skill_moves",
                "מאפיינים פיזיים תומכים."],
               ["6 תכונות הליבה", "pace, shooting, passing, dribbling, defending, physic",
                "\"הפנים\" של הכרטיס — הבסיס לדמיון ולקיבוץ."],
               ["תכונות מפורטות", "finishing, vision, sprint_speed, standing_tackle …",
                "וקטור סגנון עשיר יותר לחישוב הדמיון."]],
              col_widths=[1.4, 2.8, 2.0])

    add_par(doc, "מה הסרנו (לא רלוונטי): 26 עמודות דירוג לפי עמדה (ls, st … gk), כל "
                 "עמודות השוער (goalkeeping_*), כתובות URL, מספרי חולצה, תאריכי חוזה "
                 "ותגיות.", bold=False)

    add_heading(doc, "המניפולציות שביצענו על טבלת השחקנים", level=2)
    add_par(doc, "לא רק ניקינו — גם הוספנו עמודות מחושבות חדשות:")
    add_table(doc,
              ["העמודה החדשה", "מה היא", "איך חושבה"],
              [["position_group", "קיבוץ העמדה ל-4 קבוצות (חלוץ/קשר/הגנה/שוער)",
                "לפי העמדה הראשונה ברשימת העמדות של השחקן."],
               ["clean_name", "שם מנורמל לצורך התאמה למאגר האירועים",
                "אותיות קטנות, הסרת סימנים דיאקריטיים (Petrić ← petric)."],
               ["ability_score", "ציון יכולת 0–100 משוקלל לפי עמדה",
                "שקלול 6 תכונות הליבה במשקלים שונים לכל עמדה."],
               ["potential_growth", "כמה השחקן צפוי להשתפר",
                "potential פחות overall."],
               ["market_efficiency_score", "כמה השחקן \"זול\" יחסית ליכולתו",
                "אחוזון יכולת פחות אחוזון שווי (על שווי בלוג)."]],
              col_widths=[1.7, 2.3, 2.2])

    add_par(doc, "תובנה חשובה מהבדיקה: ל-11% מהשחקנים חסרות 6 תכונות הליבה. זו אינה "
                 "תקלת נתונים — אלה בדיוק 2,045 השוערים, שלהם אין במאגר את תכונות "
                 "שחקני השדה. לכן השוערים מטופלים בנפרד (ולכן גם ניתוחי הדמיון "
                 "מתבצעים כרגע על שחקני שדה בלבד).", bold=True)

    # ----- 3. טבלת האירועים -----
    add_heading(doc, "3. טבלת האירועים — כאן הייתה העבודה הקשה", level=1)
    add_par(doc, "זו הטבלה שדרשה את רוב העבודה המקדימה. כל שורה היא אירוע בודד "
                 "במשחק (בעיטה, גול, כרטיס, קרן, עבירה...). הבעיה: כמעט הכול מקודד "
                 "כמספרים, לא כטקסט. לדוגמה, event_type=1 פירושו \"ניסיון/בעיטה\", "
                 "event_type=4 פירושו \"כרטיס צהוב\".")
    add_par(doc, "כך נראתה שורה גולמית (מקוצר): מספרי קודים בלבד —", bold=True)
    add_table(doc,
              ["event_type", "is_goal", "bodypart", "location", "player"],
              [["1", "1", "2", "3", "mladen petric"],
               ["3", "0", "—", "—", "sergio ramos"],
               ["4", "0", "—", "—", "—"]],
              col_widths=[1.3, 1.0, 1.2, 1.1, 1.9])
    add_par(doc, "כלומר: השורה הראשונה היא בעיטה (1) שהיא גם גול (is_goal=1), ברגל "
                 "שמאל (2), ממרכז הרחבה (3). בלי מילון פענוח — אי אפשר להבין כלום.")

    add_heading(doc, "3א. תרגום הקודים לשמות", level=2)
    add_par(doc, "בנינו מילון פענוח (שאומת מול הערכים האמיתיים בקובץ) שמתרגם כל קוד "
                 "לשם קריא. דוגמה לחלק מ-event_type:")
    add_table(doc,
              ["קוד", "משמעות", "קוד", "משמעות"],
              [["1", "ניסיון (בעיטה)", "6", "כרטיס אדום"],
               ["2", "קרן", "8", "עבירה שהושגה"],
               ["3", "עבירה", "9", "נבדל"],
               ["4", "כרטיס צהוב", "10", "יד"],
               ["5", "צהוב שני", "11", "פנדל שניתן"]],
              col_widths=[0.8, 2.3, 0.8, 2.3])

    add_heading(doc, "3ב. בניית 12 עמודות בינאריות (0/1)", level=2)
    add_par(doc, "מתוך הקודים יצרנו עמודות \"דגל\" פשוטות שאפשר לסכום בהמשך. כל אחת "
                 "מקבלת 1 אם האירוע מתאים, אחרת 0:")
    add_table(doc,
              ["העמודה", "מתי = 1", "לשם מה"],
              [["is_shot", "event_type == 1", "ספירת בעיטות"],
               ["is_goal", "כבר קיים במקור", "ספירת גולים"],
               ["is_key_pass", "event_type2 == 12", "מסירות מפתח (≈ בישולים)"],
               ["is_box_shot", "מיקום בתוך הרחבה", "איכות ההזדמנות"],
               ["is_left_foot / is_right_foot", "רגל שמאל / ימין", "בסיס לאיזון דו-רגלי"],
               ["is_header", "נגיחה", "ניתוח סוג הבעיטה"],
               ["is_on_target", "shot_outcome == 1", "דיוק בעיטות"],
               ["is_yellow / is_red", "כרטיסים", "ציון משמעת"],
               ["is_foul", "event_type == 3", "ספירת עבירות"],
               ["is_through_ball_assist", "assist_method == 4", "יצירתיות"]],
              col_widths=[2.1, 2.0, 2.1])

    add_par(doc, "סף יחיד שהגדרנו: \"בעיטה מתוך הרחבה\" (is_box_shot) — 7 קודי המיקום "
                 "שנמצאים גאומטרית בתוך רחבת ה-16, כי בעיטה מתוך הרחבה היא הזדמנות "
                 "איכותית.", bold=False)

    add_heading(doc, "3ג. נרמול שמות לצורך התאמה (clean_name)", level=2)
    add_par(doc, "כדי לחבר אירוע לשחקן הנכון, נרמלנו את שם השחקן (אותיות קטנות, "
                 "הסרת סימנים מיוחדים). שורות של אירועים קבוצתיים בלי שחקן (למשל "
                 "חלק מהקרנות) נשארות בלי שם — וזה לגיטימי.")

    add_heading(doc, "3ד. למה יש הרבה תאים ריקים — וזה תקין", level=2)
    add_par(doc, "ב-75% ומעלה מהעמודות יש ערכים חסרים. זו אינה שחיתות נתונים אלא "
                 "מבנה: עמודות כמו shot_place או bodypart רלוונטיות רק לבעיטות, ולכן "
                 "ריקות בעבירות/קרנות. כשמסכמים לרמת שחקן — ריק פשוט נספר כ-0, בלי "
                 "צורך להמציא ערכים (בלי impute).", bold=True)

    # ----- 4. האגרגציות -----
    add_heading(doc, "4. מהאירועים אל השחקן — שתי אגרגציות", level=1)
    add_par(doc, "האירועים הבודדים לא שימושיים ישירות. כיווצנו אותם בשני שלבים:")
    add_table(doc,
              ["השלב", "המפתח", "התוצאה", "כמה שורות"],
              [["שלב 5: שחקן × משחק", "clean_name + מזהה משחק",
                "סכום הדגלים בכל משחק (בעיטות, גולים, כרטיסים...)", "~228,000"],
               ["שלב 6: שחקן בודד", "clean_name",
                "סיכום עונתי + ממוצעים + שיעורים + 4 ציונים", "6,106"]],
              col_widths=[1.7, 1.7, 2.4, 0.9])
    add_par(doc, "טבלת \"שחקן × משחק\" היא גם הבסיס לספירת \"דאבלים\" "
                 "(matches_with_2_plus_goals) — משחקים שבהם שחקן הבקיע 2+ גולים.")

    add_heading(doc, "4א. ארבעת הציונים המחושבים (שלב 6)", level=2)
    add_par(doc, "כאן הפכנו את הספירות הגולמיות לציונים שמספרים סיפור על השחקן:")
    add_table(doc,
              ["הציון", "מה הוא מודד", "עיקרון החישוב"],
              [["attacking_involvement_score", "מעורבות התקפית (גולים מובילים)",
                "שקלול גולים/מסירות/בעיטות → אחוזון 0–100."],
               ["creative_score", "יצירתיות ובישולים",
                "דגש על מסירות עומק ומסירות מפתח → אחוזון."],
               ["discipline_score", "משמעת (מעט כרטיסים)",
                "100 = אף כרטיס; יורד לפי כרטיסים למשחק (אדום שווה 3)."],
               ["foot_balance_score", "איזון דו-רגלי",
                "יחס הרגל החלשה לחזקה לפי בעיטות (100 = מאוזן)."]],
              col_widths=[2.2, 2.1, 2.0])
    add_par(doc, "כדי שהציונים יהיו אמינים, חישבנו אותם רק לשחקנים עם 5+ משחקים "
                 "(אחרת ממוצע מ-1–2 משחקים הוא רעש). בדיקת שפיות: בראש הרשימות יצאו "
                 "שמות צפויים כמו Ronaldo ו-Messi.", bold=False)

    # ----- 5. המיזוג -----
    add_heading(doc, "5. המיזוג — הטבלה המרכזית", level=1)
    add_par(doc, "מיזגנו את פרופיל FC24 עם ביצועי האירועים לפי clean_name (חיבור "
                 "שמאלי: כל שחקני FC24 נשמרים). שחקן שנמצאה לו התאמה מקבל "
                 "has_event_data = True.")
    add_par(doc, "פער השנים (לב ה-PoC): FC24 הוא מ-2023 והאירועים מ-2012–2017, ולכן "
                 "רק 902 שחקנים (4.9%) קיימים בשני המקורות. את הפער הזה לא הסתרנו — "
                 "סימנו אותו במפורש. התוצאה: טבלה מרכזית של 18,350 שחקנים × 78 "
                 "עמודות.", bold=True)
    add_par(doc, "המשמעות המעשית: חיפוש לפי פרופיל עובד על כל 18,350 השחקנים; ניתוחי "
                 "הביצועים האמיתיים מתבצעים על 902 השחקנים שיש להם נתוני מגרש.")

    # ----- 6. סיכום הצינור -----
    add_heading(doc, "6. סיכום — 7 הטבלאות בצינור", level=1)
    add_table(doc,
              ["#", "הטבלה", "רמת השורה", "תפקיד"],
              [["1", "male_players (גולמי)", "שחקן × גרסה", "מקור פרופיל FC24"],
               ["2", "events (גולמי)", "אירוע", "מקור ביצועים"],
               ["3", "clean_players", "שחקן", "ניקוי + ציונים מחושבים"],
               ["4", "clean_events", "אירוע", "תרגום קודים + 12 דגלים"],
               ["5", "player_match_stats", "שחקן × משחק", "אגרגציה + בסיס לדאבלים"],
               ["6", "player_event_stats", "שחקן", "סיכום + 4 ציונים"],
               ["7", "final_scouting_table", "שחקן", "הטבלה המרכזית (מיזוג)"]],
              col_widths=[0.4, 2.0, 1.5, 2.3])

    add_par(doc, "בקצרה: התחלנו משתי טבלאות גולמיות, ניקינו אותן, תרגמנו את שפת "
                 "הקודים של האירועים, בנינו דגלים וציונים, כיווצנו לרמת שחקן, "
                 "ומיזגנו להכול לטבלה אחת מרכזית שעליה רץ כל הסוכן.", bold=True,
            color=AZURE)

    add_par(doc, "— נכתב עבור פרויקט הגמר Master Scout, 2026.", size=9.5,
            color=RGBColor(0x70, 0x70, 0x70))

    out = (Path(__file__).resolve().parent.parent /
           "הסבר מורחב על טבלאות של פרויקט סקואטינג כדורגל פאדי 2026.docx")
    doc.save(out)
    # הדפסה בטוחה לקונסולה (שמות עבריים עלולים להיכשל בקידוד cp1252 בחלונות)
    try:
        print("SAVED:", out)
    except UnicodeEncodeError:
        print("SAVED:", out.name.encode("ascii", "replace").decode("ascii"))


if __name__ == "__main__":
    build()
