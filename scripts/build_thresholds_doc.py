"""
build_thresholds_doc.py — generates the central thresholds document (Word, Hebrew/RTL).

This is a LIVING document: every stage that introduces a threshold/weight adds a
section here, with the chosen value AND the reasoning. Player/team names stay in
English; explanations are in Hebrew.

Run:  conda activate masterscout && python scripts/build_thresholds_doc.py
Out:  docs/ספים_והחלטות_Master_Scout.docx
"""

# ייבוא Path לעבודה עם נתיבי קבצים
from pathlib import Path
# ייבוא Document ליצירת מסמך Word
from docx import Document
# ייבוא יחידות עיצוב (גודל, צבע, אינצ'ים)
from docx.shared import Pt, RGBColor, Inches
# ייבוא קבוע יישור פסקה
from docx.enum.text import WD_ALIGN_PARAGRAPH
# ייבוא qn לגישה לרכיבי XML של docx
from docx.oxml.ns import qn
# ייבוא OxmlElement ליצירת רכיבי XML (RTL)
from docx.oxml import OxmlElement

# שורש הפרויקט
ROOT = Path(__file__).resolve().parent.parent
# נתיב מסמך הפלט
OUT = ROOT / "docs" / "ספים_והחלטות_Master_Scout.docx"


# ---------- RTL helpers ----------
# פונקציית עזר: הופכת פסקה לכיווניות ימין-לשמאל
def _set_rtl(paragraph):
    # מקבלים/יוצרים את מאפייני הפסקה
    pPr = paragraph._p.get_or_add_pPr()
    # יוצרים רכיב bidi
    bidi = OxmlElement("w:bidi")
    # מציבים את ערכו ל-1 (RTL פעיל)
    bidi.set(qn("w:val"), "1")
    # מוסיפים אותו לפסקה
    pPr.append(bidi)
    # מיישרים את הפסקה לימין
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# פונקציית עזר: הופכת טבלה שלמה ל-RTL
def _rtl_table(table):
    # מאפייני הטבלה
    tblPr = table._tbl.tblPr
    # רכיב bidi חזותי לטבלה
    bidi = OxmlElement("w:bidiVisual")
    # מוסיפים אותו
    tblPr.append(bidi)
    # עוברים על כל שורה
    for row in table.rows:
        # ועל כל תא
        for cell in row.cells:
            # ועל כל פסקה בתא — מגדירים RTL
            for p in cell.paragraphs:
                _set_rtl(p)


# פונקציית עזר: מוסיפה כותרת RTL ברמה נתונה
def h(doc, text, level=1):
    # מוסיפים כותרת
    p = doc.add_heading(text, level=level)
    # הופכים אותה ל-RTL
    _set_rtl(p)
    # מחזירים את הפסקה
    return p


# פונקציית עזר: מוסיפה פסקת טקסט RTL
def para(doc, text, bold=False, size=11):
    # מוסיפים פסקה
    p = doc.add_paragraph()
    # מוסיפים את הטקסט כריצה
    run = p.add_run(text)
    # מדגישים אם התבקש
    run.bold = bold
    # מגדירים גודל גופן
    run.font.size = Pt(size)
    # הופכים ל-RTL
    _set_rtl(p)
    # מחזירים את הפסקה
    return p


# פונקציית עזר: מוסיפה פריט רשימה (תבליט) RTL
def bullet(doc, text):
    # מוסיפים פסקה בסגנון רשימת תבליטים
    p = doc.add_paragraph(style="List Bullet")
    # מוסיפים את הטקסט
    p.add_run(text)
    # הופכים ל-RTL
    _set_rtl(p)
    # מחזירים את הפסקה
    return p


# פונקציית עזר: מטמיעה תמונה (אם קיימת) במרכז, עם כיתוב אופציונלי
def image(doc, path, width_in=5.8, caption=None):
    """Embed an image (if it exists), centered, with an optional caption."""
    # ממירים לנתיב
    path = Path(path)
    # אם הקובץ לא קיים — לא מטמיעים
    if not path.exists():
        return
    # מוסיפים פסקה ממורכזת
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # מטמיעים את התמונה ברוחב נתון
    p.add_run().add_picture(str(path), width=Inches(width_in))
    # אם יש כיתוב — מוסיפים אותו ממורכז ונטוי
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


# פונקציית עזר: מוסיפה טבלה עם כותרות מודגשות, RTL
def table(doc, headers, rows):
    # יוצרים טבלה עם שורת כותרת
    t = doc.add_table(rows=1, cols=len(headers))
    # מגדירים סגנון טבלה
    t.style = "Light Grid Accent 1"
    # ממלאים את כותרות העמודות (מודגשות)
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].add_run(htext).bold = True
    # ממלאים את שורות הנתונים
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].paragraphs[0].add_run(str(val))
    # הופכים את הטבלה ל-RTL
    _rtl_table(t)
    # מחזירים אותה
    return t


# ====================================================================
# בונה את מסמך הספים וההחלטות המלא
def build():
    # יוצרים מסמך חדש
    doc = Document()
    # default font
    # גופן ברירת מחדל
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # כותרת ראשית
    title = doc.add_heading("מסמך ספים והחלטות — פרויקט Master Scout", level=0)
    # RTL לכותרת
    _set_rtl(title)
    # כותרת משנה ותיאור המסמך
    para(doc, "סוכן סקאוטינג שחקני כדורגל · סדנת AI & ML · קורס 277302",
         bold=True)
    # פסקת פתיחה המסבירה את מטרת המסמך
    para(doc, "המסמך מרכז את כל הספים, המשקלים והסיפים שנקבעו בפרויקט — הערך שנבחר "
              "וההנמקה מאחוריו. כל סף מתועד גם כהערה בקוד עצמו. שמות שחקנים/קבוצות "
              "באנגלית; ההסברים בעברית. זהו מסמך חי — מתעדכן בכל שלב שמוסיף סף.")

    # ---------------- Stage 1 ----------------
    # שלב 1 — מבנה הנתונים ומיפוי הקודים
    h(doc, "שלב 1 — מבנה הנתונים ומיפוי קודים", 1)
    para(doc, "אין ספים מספריים. נבחרו עמודות רלוונטיות מכל קובץ ונבנתה טבלת מיפוי "
              "לקודי האירועים (event_type, location, bodypart וכו'). פירוט מלא: "
              "docs/01_schema.md.")
    bullet(doc, "החלטה: מתוך male_players.csv (10 גרסאות FIFA) נשמרת רק "
                "fifa_version==24 — הפרופיל העדכני ביותר, שורה אחת לכל שחקן.")
    bullet(doc, "נימוק: הקובץ מכיל היסטוריה של 10 שנות FIFA; לסקאוטינג רלוונטי "
                "המצב הנוכחי בלבד (18,350 שחקנים ב-FC24).")

    # ---------------- Stage 3 ----------------
    # שלב 3 — ניקוי השחקנים והציונים המחושבים
    h(doc, "שלב 3 — ניקוי שחקנים וציונים מחושבים (clean_players)", 1)

    h(doc, "3.1 — קיבוץ עמדות (position_group)", 2)
    para(doc, "כל שחקן מסווג לאחת מארבע קבוצות לפי העמדה הראשונה ב-player_positions.")
    para(doc, "נימוק: העמדה הראשונה היא העמדה העיקרית של השחקן; הקיבוץ מאפשר "
              "השוואה והמלצה לפי תפקיד (חלוץ מול חלוץ, מגן מול מגן).")
    table(doc, ["קבוצה", "עמדות FC כלולות"], [
        ["GK", "GK"],
        ["Defender", "CB, RB, LB, RWB, LWB"],
        ["Midfielder", "CDM, CM, CAM, RM, LM"],
        ["Forward", "RW, LW, CF, ST"],
    ])

    h(doc, "3.2 — ציון יכולת (ability_score), טווח 0–100", 2)
    para(doc, "ממוצע משוקלל של 6 הסטטים הראשיים (pace, shooting, passing, dribbling, "
              "defending, physic), כאשר המשקלים שונים לכל קבוצת עמדה.")
    para(doc, "נימוק: כל תפקיד נמדד לפי התכונות שחשובות לו — חלוץ לפי בעיטה ומהירות, "
              "מגן לפי הגנה וכוח, קשר לפי מסירה ודריבל. המשקלים מסתכמים ל-1.0.")
    table(doc, ["קבוצה", "המשקלים שנבחרו"], [
        ["Forward", "shooting 0.30 · pace 0.20 · dribbling 0.20 · passing 0.10 · physic 0.10 · defending 0.10"],
        ["Midfielder", "passing 0.30 · dribbling 0.25 · pace 0.15 · shooting 0.10 · defending 0.10 · physic 0.10"],
        ["Defender", "defending 0.40 · physic 0.25 · pace 0.15 · passing 0.10 · dribbling 0.05 · shooting 0.05"],
        ["GK", "אין סטטים התקפיים → נעשה שימוש ב-overall כפי שהוא"],
    ])
    bullet(doc, "טיפול בערכים חסרים: אם סטט בודד חסר, המשקל מנורמל מחדש על הסטטים "
                "הקיימים (לא נספר כאפס).")
    bullet(doc, "שוערים: 2,045 שוערים חסרים את 6 הסטטים (ראה docs/02_eda.md), לכן "
                "עבורם ה-ability_score = overall.")

    h(doc, "3.3 — ציון יעילות שוק (market_efficiency_score), טווח −100..100", 2)
    para(doc, "ההפרש בין אחוזון היכולת של השחקן לאחוזון השווי שלו. ערך חיובי = יותר "
              "יכולת ממה שהמחיר מרמז (מועמד למציאה).")
    para(doc, "נימוק: 'מציאה' = יחס יכולת/מחיר גבוה. value_eur עובר טרנספורמציית log "
              "לפני הדירוג כי הוא מוטה ימינה בקיצוניות (ממוצע 2.8M מול חציון 1.0M).")
    bullet(doc, "MIN_VALUE_EUR = 10,000 € — שחקנים מתחת לסף (סוכנים חופשיים / חוסר "
                "נתונים) מקבלים NaN, כי חישוב יעילות עבורם חסר משמעות.")
    bullet(doc, "הערה: זהו סיגנל גולמי בלבד. דגל ה'מציאה' האמיתי (המשלב גיל ופוטנציאל) "
                "מחושב בשלב 11 עם Isolation Forest. בשלב זה המציאות המובילות הן ותיקים "
                "זולים — נכון מתמטית, אך לא 'תכשיט' — ולכן נדרש ההקשר של שלב 11.")

    h(doc, "3.4 — מפתח התאמת שמות (clean_name)", 2)
    para(doc, "נורמליזציה של long_name: אותיות קטנות, הסרת סימנים דיאקריטיים "
              "(Petrić→petric), הסרת פיסוק, איחוד רווחים. משמש רק לחיבור עם טבלת "
              "האירועים בשלב 7; שם התצוגה נשאר באנגלית המקורית.")
    bullet(doc, "מגבלה ידועה: ב-events השמות קצרים (firstname lastname) בעוד "
                "long_name כולל שמות אמצעיים — חלק מהכוכבים לא יתאימו. מתועד כמגבלה "
                "(פרק 9 במחוון).")

    # ---------------- Stage 4 ----------------
    # שלב 4 — ניקוי האירועים ומיפוי הקודים
    h(doc, "שלב 4 — ניקוי אירועים ומיפוי קודים (clean_events)", 1)
    para(doc, "כל קודי האירועים תורגמו לעמודות טקסט (event_type_name, bodypart_name "
              "וכו') ונבנו עמודות בינאריות לכל מאפיין רלוונטי. המיפוי המלא: "
              "docs/01_schema.md.")

    h(doc, "4.1 — סף 'בעיטה בתוך הרחבה' (is_box_shot)", 2)
    para(doc, "בעיטה מסומנת is_box_shot=1 אם קוד ה-location שלה נמצא בקבוצת הקודים "
              "שמייצגים מיקום בתוך רחבת ה-16.")
    para(doc, "נימוק: בעיטה מתוך הרחבה היא הזדמנות איכותית; היחס בין בעיטות-רחבה לכלל "
              "הבעיטות מאפיין חלוצי רחבה לעומת בועטים מטווח רחוק.")
    table(doc, ["קוד location", "מיקום (בתוך הרחבה)"], [
        ["3", "Centre of the box"],
        ["9", "Left side of the box"],
        ["10", "Left side of the six yard box"],
        ["11", "Right side of the box"],
        ["12", "Right side of the six yard box"],
        ["13", "Very close range"],
        ["14", "Penalty spot"],
    ])

    h(doc, "4.2 — עמודות בינאריות שנגזרו", 2)
    para(doc, "כל עמודה היא 0/1 ומשמשת לאגרגציה ברמת שחקן (שלבים 5–6):")
    bullet(doc, "is_shot (event_type=1), is_goal (קיים במקור), is_key_pass "
                "(event_type2=12), is_box_shot (ראה 4.1).")
    bullet(doc, "is_left_foot / is_right_foot / is_header (bodypart=2/1/3) — בסיס "
                "לציון איזון הרגליים (foot balance).")
    bullet(doc, "is_on_target (shot_outcome=1), is_yellow (event_type=4), is_red "
                "(event_type∈{5,6}), is_foul (event_type=3) — בסיס לציון משמעת.")
    bullet(doc, "is_through_ball_assist (assist_method=4) — בסיס לציון יצירתיות.")
    bullet(doc, "החלטה: אירועים ללא player (אירועי קבוצה כמו קרנות) נשמרים, אך "
                "clean_name שלהם NaN ולכן יסוננו באגרגציה לשחקן.")

    # ---------------- Stage 6 ----------------
    # שלב 6 — אגרגציה לשחקן וציוני הביצוע
    h(doc, "שלב 6 — אגרגציה לשחקן וציוני ביצוע (player_event_stats)", 1)
    para(doc, "שורה אחת לכל שחקן עם סכומים, ערכים per_match, יחסים (0–1) וארבעה "
              "ציונים מחושבים. כאן נדרשים ספי מהימנות כדי למנוע רעש ממדגם קטן.")

    h(doc, "6.1 — ספי מהימנות", 2)
    bullet(doc, "MIN_MATCHES = 5 — מתחת לסף זה היחסים/הציונים אינם אמינים (שחקן עם "
                "משחק אחד וגול אחד היה מקבל goals_per_match=1.0). שחקנים מתחת לסף "
                "מקבלים NaN בציוני האחוזון.")
    bullet(doc, "MIN_FOOT_SHOTS = 5 — אי אפשר להעריך דו-רגליות מ-1–2 בעיטות; מתחת "
                "לסף foot_balance_score = NaN.")
    para(doc, "מגבלה ידועה: שחקנים בדיוק מעל הסף (כ-5–6 משחקים) עלולים להגיע "
              "לאחוזונים גבוהים במדגם קטן. מתועד כמגבלה; הסוכן יכול להחיל סינון "
              "מינימום-משחקים מחמיר יותר.")

    h(doc, "6.2 — ציון מעורבות התקפית (attacking_involvement_score), 0–100", 2)
    para(doc, "דירוג אחוזוני של שילוב משוקלל per-match (גולים במשקל הגבוה ביותר), "
              "בקרב שחקנים עם matches ≥ MIN_MATCHES.")
    table(doc, ["רכיב (per match)", "משקל"], [
        ["goals_per_match", "3.0"],
        ["key_passes_per_match", "2.0"],
        ["box_shots_per_match", "1.0"],
        ["shots_per_match", "0.5"],
    ])
    para(doc, "ולידציה: המובילים הם Cristiano Ronaldo, Lionel Messi, Sergio Agüero, "
              "Luis Suárez, Zlatan Ibrahimović — חלוצי עילית אמיתיים.")

    h(doc, "6.3 — ציון יצירתיות (creative_score), 0–100", 2)
    para(doc, "דירוג אחוזוני של שילוב per-match; מסירות פתיחה (through ball) נדירות "
              "ויצירתיות יותר ולכן משקל גבוה יותר.")
    table(doc, ["רכיב (per match)", "משקל"], [
        ["through_ball_assists_per_match", "3.0"],
        ["key_passes_per_match", "1.0"],
    ])

    h(doc, "6.4 — ציון משמעת (discipline_score), 0–100", 2)
    para(doc, "100 = מעולם לא הוזהר; 0 = ≥ DISCIPLINE_CARD_CAP כרטיסים משוקללים "
              "למשחק. כרטיס אדום שווה ל-RED_CARD_WEIGHT כרטיסים צהובים.")
    bullet(doc, "DISCIPLINE_CARD_CAP = 0.5 כרטיסים משוקללים למשחק = הציון הגרוע ביותר.")
    bullet(doc, "RED_CARD_WEIGHT = 3 — אדום חמור פי 3 מצהוב.")
    para(doc, "נימוק: 0.5 כרטיסים משוקללים למשחק (למשל צהוב כל 2 משחקים) הוא רף "
              "סביר לשחקן 'בעייתי' מבחינת משמעת.")

    h(doc, "6.5 — ציון איזון רגליים (foot_balance_score), 0–100", 2)
    para(doc, "100 = איזון מושלם בין בעיטות רגל ימין לשמאל; 0 = חד-רגלי מוחלט. "
              "מחושב כ-min(left,right)/max(left,right)×100, בכפוף ל-MIN_FOOT_SHOTS.")

    # ---------------- Stage 7 ----------------
    # שלב 7 — בניית הטבלה המרכזית (המיזוג)
    h(doc, "שלב 7 — הטבלה המרכזית (final_scouting_table)", 1)
    para(doc, "מיזוג פרופיל FC24 (clean_players) עם מדדי הביצוע (player_event_stats) "
              "לטבלה אחת שעליה רץ כל הסוכן.")
    h(doc, "7.1 — החלטת המיזוג", 2)
    bullet(doc, "סוג חיבור: LEFT JOIN (players ← event_stats) על clean_name. כל "
                "שחקני FC24 נשמרים; מי שאין לו התאמת אירועים מקבל has_event_data=False.")
    bullet(doc, "נימוק: אסור 'לאבד' שחקנים בגלל פער הכיסוי — מסמנים את מקור הנתונים "
                "במפורש (data_source_note) במקום למחוק.")
    h(doc, "7.2 — היקף החפיפה (ממצא מרכזי ל-PoC)", 2)
    table(doc, ["מצב", "כמות", "אחוז"], [
        ["סך שחקני FC24", "18,350", "100%"],
        ["עם נתוני אירועים (has_event_data=True)", "902", "4.9%"],
        ["FC24 בלבד", "17,448", "95.1%"],
    ])
    para(doc, "מתוך 6,106 שחקני ה-events, רק 902 קיימים גם ב-FC24 2023 — היתר פרשו/"
              "עזבו עד 2023. זהו פער השנים (events 2012–2017 מול FC24 2023) ומהווה את "
              "המגבלה המרכזית (פרק 9 במחוון). 902 שחקנים מספיקים בשפע להדגמת דמיון/"
              "קלאסטרינג/חריגות; חיפוש לפי פרופיל עובד על כל 18,350.")
    bullet(doc, "מגבלה נוספת: 40 התנגשויות שם (שחקני FC24 שונים עם אותו clean_name) — "
                "המיזוג עלול לשייך להם אותם נתוני אירועים. מתועד.")

    # ---------------- Stage 8 ----------------
    # שלב 8 — פונקציות החיפוש הפרמטריות
    h(doc, "שלב 8 — פונקציות חיפוש (search.py)", 1)
    para(doc, "עקרון מנחה: כל סף הוא פרמטר עם ערך ברירת מחדל, כך שקל לשנותו. "
              "חיפושי פרופיל רצים על כל 18,350 השחקנים; חיפושי ביצוע "
              "(התקפי/יצירתי/משמעת/דו-רגלי/ברייסים) דורשים has_event_data=True "
              "(כ-902 שחקנים).")
    h(doc, "8.1 — ערכי ברירת המחדל של הספים", 2)
    table(doc, ["פרמטר", "ברירת מחדל", "משמעות / נימוק"], [
        ["top_n", "20", "מספר התוצאות המוחזרות"],
        ["min_discipline (defenders)", "70", "ציון משמעת מינימלי; 70 = מעט מאוד כרטיסים"],
        ["min_foot_balance (two-footed)", "60", "איזון רגליים מינימלי לסיווג כדו-רגלי"],
        ["min_total_shots (two-footed)", "20", "סף נפח — מונע 'איזון מושלם' ממדגם קטן"],
        ["min_braces", "2", "מינימום משחקי 2+ גולים"],
    ])
    para(doc, "כל שאר הספים (max_age, min_pace, max_value_eur, min_overall, "
              "min_goals_per_match וכו') הם פרמטרים אופציונליים ללא ברירת מחדל — "
              "מופעלים רק אם הסוכן מחלץ אותם מהשאלה.")
    para(doc, "ולידציה: חיפוש דו-רגליים מחזיר את Ousmane Dembélé (דו-רגלי ידוע); "
              "חיפוש ברייסים מחזיר את Lewandowski (35 משחקי 2+ גולים).")

    # ---------------- Stage 9 ----------------
    # שלב 9 — דמיון שחקנים ב-Cosine (שימוש #1)
    h(doc, "שלב 9 — דמיון שחקנים ב-Cosine (similarity.py) — שימוש #1", 1)
    para(doc, "מציאת שחקנים דומים: נרמול פיצ'רים + Cosine. זהו האנלוג המספרי של "
              "TF-IDF+Cosine מהקורס — במקום משקלי TF-IDF, וקטור תכונות מנורמל.")
    h(doc, "9.1 — בחירת פיצ'רים", 2)
    para(doc, "14 תכונות FC24 שפורשׂות סגנון משחק: pace, shooting, passing, "
              "dribbling, defending, physic, skill_ball_control, skill_dribbling, "
              "mentality_vision, attacking_finishing, power_strength, "
              "movement_sprint_speed, movement_acceleration, defending_standing_tackle. "
              "זמינות לכל שחקני השדה (לא רק 902).")
    h(doc, "9.2 — נרמול ומדד", 2)
    bullet(doc, "StandardScaler (z-score) — משווה את דפוס היכולת (מעל/מתחת לממוצע), "
                "מבחין יותר מווקטורים חיוביים גולמיים.")
    bullet(doc, "Cosine בטווח [-1,1] ממופה ל-[0,1] ע\"י (x+1)/2: 1=זהה, 0.5=לא קשור, "
                "0=הפוך. הציונים גבוהים (~0.99) כי שחקנים מאותה עמדה דומים מהותית — "
                "מתפרשים יחסית (הדירוג חשוב, לא הערך המוחלט).")
    h(doc, "9.3 — השוואת מודלים (קריטי לפרק 5+6): עם/בלי קטגוריה", 2)
    para(doc, "אותה שאילתה מורצת פעמיים: WITH category = המועמדים מוגבלים לאותה "
              "position_group; WITHOUT = כל שחקני השדה.")
    para(doc, "תוצאה לדוגמה (Kevin De Bruyne): WITH מחזיר קשרים יצירתיים בלבד "
              "(Bruno Fernandes, Mac Allister, Lo Celso); WITHOUT מכניס גם שחקנים "
              "מעמדות אחרות עם פרופיל סטטיסטי דומה. חפיפה ~5/8. מסקנה: הקטגוריה "
              "ממקדת בתפקיד — בדרך כלל שימושית יותר לסקאוטינג.")

    # ---------------- Stage 10 ----------------
    # שלב 10 — קיבוץ סגנונות משחק ב-K-Means (שימוש #1)
    h(doc, "שלב 10 — קיבוץ סגנונות משחק ב-K-Means (clustering.py) — שימוש #1", 1)
    para(doc, "קיבוץ שחקנים לסגנונות לפי אותן 14 תכונות FC24 (StandardScaler). "
              "רץ על שחקני שדה עם תכונות מלאות (16,305); שוערים מקבלים cluster_id=-1.")
    h(doc, "10.1 — בחירת מספר הקלאסטרים k (שיטת Elbow)", 2)
    bullet(doc, "k = 5 (DEFAULT_K). נבחר לפי 'מרפק' עקומת ה-inertia (k=2..10) "
                "סביב 4–5, בשילוב פרשנות: 5 קלאסטרים מניבים סגנונות מובחנים וברי-שיום.")
    bullet(doc, "ה-silhouette הגבוה ביותר ב-k=2 (~0.28) ויורד — תופעה רגילה בנתוני "
                "שחקנים (סגנונות רציפים וחופפים). העדפנו פרשנות על פני silhouette "
                "מקסימלי. הגרף נשמר ב-reports/figures/10_kmeans_elbow.png.")
    bullet(doc, "RANDOM_STATE = 42, n_init = 10 — לשחזוריות.")
    h(doc, "10.2 — הסגנונות שהתקבלו (ולידציה)", 2)
    table(doc, ["קלאסטר", "תווית", "דוגמאות"], [
        ["2", "technical / playmaking", "Mbappé, De Bruyne, Haaland, Messi, Benzema"],
        ["4", "playmaking / defensive", "Rodri, Casemiro, Kroos, Rúben Dias"],
        ["3", "defensive / physical", "בלמים פיזיים (92% מגנים)"],
        ["0", "defensive / fast", "Varane, Rüdiger, Araujo"],
        ["1", "goal-scoring / fast", "חלוצים מהירים (defending נמוך)"],
    ])
    para(doc, "התוויות נוצרות אוטומטית מהתכונות הבולטות (2 התכונות הכי מעל הממוצע). "
              "הקלאסטרים תואמים אינטואיציה כדורגלית — ולידציה חזקה.")

    image(doc, ROOT / "reports/figures/10_kmeans_elbow.png",
          caption="גרף Elbow + Silhouette — בחירת k=5")
    image(doc, ROOT / "reports/figures/10_kmeans_clusters.png",
          caption="הקלאסטרים בהיטל PCA דו-ממדי (צבע = סגנון משחק)")

    # ---------------- Stage 11 ----------------
    # שלב 11 — גילוי חריגות ומציאות (שימוש #2)
    h(doc, "שלב 11 — גילוי חריגות (anomaly.py) — שימוש #2", 1)
    para(doc, "שני שימושים: (א) איתור מציאות (יכולת גבוהה במחיר נמוך); (ב) חריגות "
              "פרופיל-מול-ביצוע (דירוג FC24 לא תואם תפוקה אמיתית). אלגוריתמים מוכנים "
              "מ-scikit-learn — אנו משתמשים, לא מאמנים.")
    h(doc, "11.1 — ספים", 2)
    bullet(doc, "ANOMALY_CONTAMINATION = 0.02 — שיעור החריגות הצפוי (Isolation "
                "Forest ו-nu של One-Class SVM). מציאות/חריגות נדירות; 2% שומר "
                "רשימה קצרה ואיכותית לסקאוט.")
    bullet(doc, "DBSCAN: eps = 0.6, min_samples = 6. נתוני השחקנים הם 'ענן רציף' "
                "בלי פערי צפיפות, ולכן גם eps קטן משאיר רק כ-14 נקודות כרעש.")
    h(doc, "11.2 — הגדרת 'מציאה'", 2)
    para(doc, "Isolation Forest על [overall, potential, age, log_value] (שחקני שדה); "
              "מתוך החריגות נשמרות המתומחרות-בחסר (market_efficiency_score>0) וממוינות. "
              "עם min_overall=78 מתקבלות מציאות עילית-וזולות: Chiellini, Sergio Ramos, "
              "Thiago Silva, Hummels, Jordi Alba — בלמים מעולים שערכם ירד עם הגיל.")
    h(doc, "11.3 — השוואת מודלים (קריטי לפרק 5+6): IF מול DBSCAN מול One-Class SVM", 2)
    table(doc, ["אלגוריתם / חיתוך", "מספר חריגות"], [
        ["Isolation Forest", "365"],
        ["DBSCAN", "14"],
        ["One-Class SVM", "394"],
        ["IF ∩ DBSCAN", "14"],
        ["שלושתם יחד", "14"],
    ])
    para(doc, "מסקנה: IF ו-OCSVM מסמנים ~2% קבוע; DBSCAN מסמן רק 14 כי הנתונים רציפים "
              "(אין פערי צפיפות) — DBSCAN פחות מתאים כאן, ו-IF הוא הכלי הנכון. 14 "
              "הנקודות שכל השלושה הסכימו עליהן הן החריגות הכי יציבות.")

    # ---------------- placeholder for future ----------------
    # מקטע שמור לשלבים הבאים שיתווספו עם ההתקדמות
    h(doc, "שלבים הבאים (יתווספו עם ההתקדמות)", 1)
    # תיאור השלבים הבאים
    para(doc, "שלב 12 (סוכן GPT — ספי סקופ/בהירות), שלב 13 (NLG), שלב 14 (בדיקות "
              "כולל מקרה כשל), שלב 15 (פריסה). יתועדו כאן.")

    # שומרים את המסמך לקובץ
    doc.save(OUT)
    # מדפיסים את נתיב הקובץ שנשמר
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
